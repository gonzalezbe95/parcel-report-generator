import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from scraper.king import KingScraper
from scraper.kitsap import KitsapScraper
from scraper.pierce import PierceScraper
from docx.oxml import OxmlElement


def set_document_font(document, font_name="Calibri", font_size=12):
    """Set default font and size for the entire document."""
    style = document.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), font_name)


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """Utility to insert a clickable hyperlink into a Word paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color if provided
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Add underline if requested
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def generate_word_report(parcel_numbers, assessor_url):
    """
    Generate a Word report for one or more parcels in a single document.

    :param parcel_numbers: comma-separated string or list of parcel numbers
    :param assessor_url: URL of county assessor
    :return: file_stream, error_message, all_data
    """
    if isinstance(parcel_numbers, str):
        parcel_list = [p.strip() for p in parcel_numbers.split(",") if p.strip()]
    else:
        parcel_list = parcel_numbers

    if not parcel_list:
        return None, "No parcel numbers provided", None

    document = Document()
    set_document_font(document, "Calibri", 12)

    # Main heading
    heading = document.add_heading("PROPERTY SUMMARY", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.bold = True

    all_data = {}

    # Loop through parcels
    for parcel_number in parcel_list:
        url_lower = assessor_url.lower()
        if "kingcounty" in url_lower:
            scraper = KingScraper(parcel_number)
        elif "kitsap.gov" in url_lower:
            scraper = KitsapScraper(parcel_number)
        elif "piercecountywa" in url_lower:
            scraper = PierceScraper(parcel_number)
        else:
            continue  # skip unsupported counties

        try:
            data = scraper.scrape()
            all_data[parcel_number] = data

            # Parcel heading
            parcel_heading = document.add_heading(f"Parcel: {parcel_number}", level=1)
            parcel_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Fields to include (genericized)
            fields = [
                "Parcel Summary",
                "Site Address",
                "Taxpayer Name",
                "Land Acres",
                "Land Use Description",
                "Legal Description",
                "Exemptions",
                "Related Parcels",
                "Property in Foreclosure"
            ]

            for key in fields:
                if key == "Parcel Summary":
                    # Create specific URL for this parcel
                    if "kingcounty" in url_lower:
                        value = f"https://blue.kingcounty.com/Assessor/eRealProperty/Detail.aspx?ParcelNbr={parcel_number}"
                    elif "kitsap.gov" in url_lower:
                        value = f"https://psearch.kitsap.gov/pdetails/Details?parcel={parcel_number}&page=general"
                    elif "piercecountywa" in url_lower:
                        value = f"https://atip.piercecountywa.gov/app/v2/propertyDetail/{parcel_number}/summary"
                    else:
                        value = assessor_url

                    # Special handling for Parcel Summary - hyperlink
                    para = document.add_paragraph(style="List Bullet")
                    run = para.add_run(f"{key}: ")
                    run.bold = True
                    add_hyperlink(para, value, value)
                else:
                    # Regular fields
                    value = data.get(key, "N/A")
                    para = document.add_paragraph(style="List Bullet")
                    run = para.add_run(f"{key}: ")
                    run.bold = True
                    para.add_run(f"{value or 'N/A'}")

            # Apply spacing
            para.paragraph_format.space_after = Pt(3)

        except Exception as e:
            # Show error for parcels that fail
            error_para = document.add_paragraph()
            error_para.add_run(
                f"Error scraping parcel {parcel_number}: {str(e)}"
            ).bold = True
            error_para.paragraph_format.space_after = Pt(3)

    # Collect addresses for filename and docx content
    addresses = []
    for parcel in parcel_list:
        address = all_data.get(parcel, {}).get("Site Address", parcel)
        safe_address = address.replace(" ", "_")
        addresses.append(safe_address)

    # Join addresses with underscores
    combined_addresses = "_".join(addresses)

    # Save to in-memory file
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream, None, all_data
